# src/extraction/docx_extractor.py — v1
"""DOCX extractor using python-docx.

Extracts text, embedded images, and tables from Word documents.
Requires the 'python-docx' package.
See spec §4 pipeline step 1b.
"""

from __future__ import annotations

import logging
from pathlib import Path

from ayextractor.core.models import (
    DocumentStructure,
    ExtractionResult,
    ImageAnalysis,
    TableData,
)
from ayextractor.extraction.base_extractor import BaseExtractor
from ayextractor.extraction.structure_detector import detect_structure

logger = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """Extractor for Word documents (.docx)."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def extract(self, content: bytes | str | Path) -> ExtractionResult:
        """Extract text, images, and tables from DOCX document."""
        try:
            import docx
        except ImportError as e:
            raise ImportError(
                "python-docx package required for DOCX extraction: "
                "pip install python-docx"
            ) from e

        doc = self._open_document(content, docx)

        text_parts: list[str] = []
        images: list[ImageAnalysis] = []
        tables: list[TableData] = []

        # Extract text from paragraphs (preserves heading hierarchy)
        for para in doc.paragraphs:
            if para.text.strip():
                # Prefix headings with markdown-style markers for structure detection
                style_name = (para.style.name or "").lower()
                if "heading" in style_name:
                    try:
                        level = int(style_name.replace("heading", "").strip())
                    except ValueError:
                        level = 1
                    text_parts.append(f"{'#' * level} {para.text}")
                else:
                    text_parts.append(para.text)

        # Extract tables
        for i, table in enumerate(doc.tables):
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows and len(rows) >= 2:
                md = self._rows_to_markdown(rows)
                tables.append(
                    TableData(
                        id=f"tbl_{i + 1:03d}",
                        content_markdown=md,
                        source_page=None,
                        origin="structured",
                    )
                )

        # Extract embedded images (metadata only)
        img_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_count += 1
                images.append(
                    ImageAnalysis(
                        id=f"img_{img_count:03d}",
                        type="photo",
                        description=f"[Embedded DOCX image, pending analysis]",
                        source_page=None,
                    )
                )

        raw_text = "\n\n".join(text_parts)
        structure = detect_structure(raw_text)

        return ExtractionResult(
            raw_text=raw_text,
            enriched_text=raw_text,
            images=images,
            tables=tables,
            structure=structure,
            language="en",
        )

    @staticmethod
    def _open_document(content: bytes | str | Path, docx_module: object) -> object:
        """Open DOCX from various input types."""
        docx_mod = docx_module  # type: ignore[assignment]
        if isinstance(content, Path):
            return docx_mod.Document(str(content))
        if isinstance(content, str):
            p = Path(content)
            if p.exists() and p.is_file():
                return docx_mod.Document(content)
            raise FileNotFoundError(f"DOCX file not found: {content}")
        # bytes: wrap in BytesIO
        import io
        return docx_mod.Document(io.BytesIO(content))

    @staticmethod
    def _rows_to_markdown(rows: list[list[str]]) -> str:
        """Convert table rows to Markdown format."""
        if not rows:
            return ""
        max_cols = max(len(r) for r in rows)
        normalized = [r + [""] * (max_cols - len(r)) for r in rows]

        lines: list[str] = []
        header = normalized[0]
        lines.append("| " + " | ".join(cell for cell in header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in normalized[1:]:
            lines.append("| " + " | ".join(cell for cell in row) + " |")
        return "\n".join(lines)
